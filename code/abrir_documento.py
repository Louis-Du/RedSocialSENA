from docx import Document
doc = Document("documents/historias_usuario_sena_mejorada.docx")

print("Número de tablas:", len(doc.tables))
for i, table in enumerate(doc.tables):
    print(f"--- Tabla {i+1} ---")
    for row in table.rows:
        print(" | ".join(cell.text.strip() for cell in row.cells))
