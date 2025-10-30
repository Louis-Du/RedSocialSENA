from docx import Document
doc = Document("documents/historias_usuario_sena_mejorada.docx")

print("Párrafos detectados:", len(doc.paragraphs))
for p in doc.paragraphs[:5]:
    print("-", p.text.strip())

print("\nTablas detectadas:", len(doc.tables))
for table in doc.tables:
    for row in table.rows:
        print(" | ".join(cell.text.strip() for cell in row.cells))
    print("-" * 40)