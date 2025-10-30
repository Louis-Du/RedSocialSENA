# crear_issues_desde_word.py
import os
import re
from github import Github
from docx import Document
from typing import List, Dict

# ---------- CONFIG ----------
# Cambia esto por el repo donde quieres crear issues (usuario/repo)
REPO_NAME = "Louis-Du/RedSocialSENA"  # ej: "lukasdiaz/red-social-sena"

# Nombre del archivo .docx en el repo (ruta relativa)
ARCHIVO_WORD = "documents/historias_usuario_sena_mejorada.docx"

# Si quieres forzar dry-run sin tocar GitHub: define DRY_RUN=1 en variables de entorno
DRY_RUN = os.getenv("DRY_RUN", "0") in ("1", "true", "True")

# ---------- HELPERS ----------
def es_inicio_hu(text: str) -> bool:
    # Detecta "Historia de Usuario RF-01" o "HU-01:" o "HU-01 -"
    return bool(re.match(r"(Historia de Usuario\s+RF-\d+)|(HU-\d+[:\s-])", text, re.IGNORECASE))

def extraer_id_nombre_desde_linea(text: str) -> (str, str):
    # Maneja: "Historia de Usuario RF-01" (id separado) o "HU-01: Como..."
    m_hu = re.match(r"(HU-\d+)\s*[:\-\s]\s*(.*)", text, re.IGNORECASE)
    if m_hu:
        return m_hu.group(1).strip(), m_hu.group(2).strip()
    m_rf = re.match(r"Historia de Usuario\s+(RF-\d+)", text, re.IGNORECASE)
    if m_rf:
        return m_rf.group(1).strip(), ""
    # fallback: tira todo como nombre
    return "", text.strip()

def normalizar_criterio(line: str) -> str:
    # "CA: Texto..." -> "Texto..."
    return re.sub(r"^CA\s*[:\-]\s*", "", line, flags=re.IGNORECASE).strip()

# ---------- PARSER ----------
from docx import Document
import re

def extraer_historias(ruta_docx):
    """
    Extrae historias de usuario y criterios de aceptación de un documento .docx,
    incluso si los criterios están dentro de celdas de tabla en una sola cadena.
    """
    doc = Document(ruta_docx)
    historias = []
    texto_total = []

    # Leer todo el texto del documento (párrafos + tablas)
    for p in doc.paragraphs:
        if p.text.strip():
            texto_total.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                contenido = cell.text.strip()
                if contenido:
                    texto_total.append(contenido)

    # Unir todo el texto
    texto_unido = "\n".join(texto_total)

    # Dividir por cada historia (RF-xx o HU-xx)
    bloques = re.split(r"(?=(?:RF|HU)-\d{1,3}:)", texto_unido)
    for bloque in bloques:
        bloque = bloque.strip()
        if not bloque:
            continue

        match = re.match(r"((?:RF|HU)-\d{1,3}):\s*(.+)", bloque, re.DOTALL)
        if match:
            id_hu = match.group(1)
            descripcion = match.group(2).strip()

            # Buscar criterios: líneas que empiecen con '-' o '•' o contengan 'CA:'
            criterios = []
            for linea in descripcion.splitlines():
                if linea.strip().startswith("-") or linea.strip().startswith("•"):
                    criterios.append(linea.strip("-• ").strip())
                elif linea.strip().startswith("CA:"):
                    criterios.append(linea.replace("CA:", "").strip())

            # Si no hay criterios, intenta detectar el bloque "Criterios de aceptación:"
            if not criterios:
                match_criterios = re.search(
                    r"Criterios\s+de\s+aceptaci[oó]n[:：]\s*(.+)",
                    descripcion,
                    re.IGNORECASE | re.DOTALL,
                )
                if match_criterios:
                    posibles = match_criterios.group(1).split("-")
                    criterios = [c.strip() for c in posibles if c.strip()]

            if not criterios:
                criterios = ["(sin criterios listados)"]

            historias.append({
                "id": id_hu,
                "descripcion": descripcion.split("Criterios")[0].strip(),
                "criterios": criterios
            })

    return historias


# ---------- CREAR ISSUE ----------
def crear_issue(repo, h: Dict):
    title_parts = []
    if h.get("id"):
        title_parts.append(h["id"])
    # para título, intenta tomar primeras 60 chars de descripción si existe
    extra = (h.get("descripcion") or "").split("\n")[0].strip()
    if extra:
        title_parts.append(extra[:60])
    title = " - ".join(title_parts) if title_parts else "Historia sin id"

    criterios = h.get("criterios", [])
    criterios_md = "\n".join([f"- {c}" for c in criterios]) if criterios else "- (sin criterios listados)"

    body = f"""**Historia de usuario:**  
{h.get('descripcion','(sin descripción)')}

**Criterios de aceptación:**  
{criterios_md}
"""

    if DRY_RUN:
        print("------ DRY RUN (no se crea issue) ------")
        print("TÍTULO:", title)
        print("CUERPO:\n", body)
        print("---------------------------------------\n")
        return None

    issue = repo.create_issue(title=title, body=body)
    agregar_issue_a_proyecto(repo, issue, "Backlog")
    print(f"✅ Issue creado: {issue.title} -> {issue.html_url}")
    return issue

def agregar_issue_a_proyecto(repo, issue, nombre_columna="Backlog"):
    """
    Vincula el issue creado con el proyecto del repositorio y lo coloca en la columna especificada.
    """
    try:
        # obtener todos los proyectos del repo
        proyectos = repo.get_projects()
        for p in proyectos:
            if p.name.lower() == "redsocialsena":  # <- nombre exacto de tu tablero
                proyecto = p
                break
        else:
            print("⚠️ Proyecto 'RedSocialSena' no encontrado en este repo.")
            return

        # buscar columna Backlog
        columna = None
        for c in proyecto.get_columns():
            if c.name.lower() == nombre_columna.lower():
                columna = c
                break
        if not columna:
            print(f"⚠️ Columna '{nombre_columna}' no encontrada en el proyecto.")
            return

        # crear tarjeta en la columna para este issue
        columna.create_card(content_id=issue.id, content_type="Issue")
        print(f"✅ Issue agregado al proyecto '{proyecto.name}' en columna '{columna.name}'.")
    except Exception as e:
        print(f"⚠️ No se pudo agregar el issue al proyecto: {e}")


# ---------- MAIN ----------
def main():
    token = os.getenv("GITHUB_TOKEN")
    if not token:
        raise SystemExit("ERROR: no se encontró la variable de entorno GITHUB_TOKEN")

    if not os.path.isfile(ARCHIVO_WORD):
        raise SystemExit(f"ERROR: no se encontró el archivo {ARCHIVO_WORD} en la ruta actual.")

    historias = extraer_historias(ARCHIVO_WORD)
    print(f"📄 Historias encontradas: {len(historias)}")
    if not historias:
        return

    # conexión (forma moderna)
    from github import Auth
    auth = Auth.Token(token)
    g = Github(auth=auth)

    try:
        me = g.get_user()
        print("Autenticado como:", me.login)
    except Exception as e:
        raise SystemExit("ERROR: No se pudo autenticar con GitHub. Revisa el token. " + str(e))

    print("Intentando obtener repo:", REPO_NAME)
    try:
        repo = g.get_repo(REPO_NAME)
    except Exception as e:
        # imprime info diagnóstica y sale, no rompas con traceback largo
        print("ERROR: get_repo falló. Detalle:", repr(e))
        # muestra una lista corta de repos accesibles para ayudar al diagnóstico
        try:
            sample = [r.full_name for i, r in enumerate(me.get_repos()) if i < 10]
            print("Repos accesibles (muestra):", sample)
        except Exception:
            pass
        raise SystemExit("Abortando por fallo al obtener el repositorio. Revisa REPO_NAME y permisos del token.")

    # si llegamos aquí, repo está bien y procedemos
    for h in historias[:3]: # limitar a las primeras 3 para pruebas
    # Extraer el número de la HU, por ejemplo "RF-23" → 23
        numero = int("".join(filter(str.isdigit, h["id"])))

    # Saltar las primeras 22
        if numero <= 22:
            print(f"⏭️ Saltando {h['id']} (ya creada)")
            continue

    crear_issue(repo, h)

if __name__ == "__main__":
    main()
